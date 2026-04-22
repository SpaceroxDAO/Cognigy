"""
Generate a Word document of clarification questions for Ralph Lauren.
These are questions we need answered to respond accurately to sheets 2, 3, and 4.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/Adam.Boyle/Downloads/NiCE - Ralph Lauren Clarification Questions.docx"

NICE_NAVY = RGBColor(0x00, 0x29, 0x5E)
NICE_TEAL = RGBColor(0x00, 0x97, 0x9D)
GREY_TEXT = RGBColor(0x44, 0x44, 0x44)
LIGHT_BG  = "EEF3FA"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Question data ─────────────────────────────────────────────────────────────
# Format: (rfp_id, why_we_need_it, question_for_rl)
QUESTIONS = {
    "Sheet 2 — AI & Virtual Agent Assessment": [
        (
            "1.01 / 1.02 — Model Transparency",
            "Your RFP asks for documentation of AI/ML model decision logic. Our answer differs significantly depending on which models are in scope — Cognigy's NLU engine (which we can document fully) vs. the third-party LLMs used for generative responses (OpenAI, Anthropic, Azure OpenAI), where decision logic is not exposed by those providers.",
            "Which AI/ML models are you expecting documentation for — the intent classification engine, the generative LLM layer, or both? What format of documentation is required (e.g., white papers, API spec, audit logs)?"
        ),
        (
            "1.03 — Model Training & Validation",
            "It's unclear whether Ralph Lauren expects to be involved in model training themselves, or whether they want visibility into how NICE/Cognigy trains and updates models on their behalf.",
            "Do you expect to train and update NLU models directly using your own team, or do you expect NICE to manage model training with change communications to you? Or both?"
        ),
        (
            "1.04 — Bias & Fairness",
            "Bias governance requirements vary widely by regulatory jurisdiction and use case. Retail customer service and regulated industries like financial services or healthcare have very different expectations.",
            "Are there specific regulatory frameworks (EU AI Act, EEOC, CCPA, etc.) or internal AI ethics policies that your bias and fairness requirements must satisfy? Are there specific high-risk decision types (e.g., routing based on customer value tier) you want governed separately?"
        ),
        (
            "2.01 / 3.01–3.03 — NLU for Retail / IVR Self-Service",
            "The RFP asks about pre-trained NLU models for retail and service. Cognigy does not ship pre-built retail intents — intents are built per deployment. We need to understand what use cases they have in mind to answer accurately.",
            "Do you have existing intent libraries or call reason taxonomies from your current IVR or contact centre platform? If so, in what format and how many intents? What are the top 10–15 reasons customers contact Ralph Lauren today?"
        ),
        (
            "2.02 — Proactive Outbound / Consent",
            "TCPA and consent management for proactive outbound voice and SMS depends entirely on where Ralph Lauren stores customer consent and how it is managed today.",
            "Where is customer contact consent currently stored (CRM, CDP, or other system)? Which channels are you seeking to use for proactive outbound — voice, SMS, email, or all three? Do you have an existing DNC/suppression list management process?"
        ),
        (
            "3.02 — IVR Migration",
            "The RFP mentions migration tools for converting DTMF IVR flows. The effort and approach change completely depending on the incumbent platform.",
            "What is your current IVR/telephony platform (e.g., Avaya, Genesys, Cisco, Amazon Connect)? Do you have IVR flow documentation in a machine-readable format (e.g., VXML, Visio, XML)? How many IVR flows or call paths are in scope for migration?"
        ),
        (
            "5.05 — Ralph Lauren NLU Governance",
            "This question asks about Ralph Lauren's ability to train and govern NLU models directly. We need to know their expected operating model — will they have a dedicated team managing this or delegate to NICE professional services?",
            "Who within Ralph Lauren would own NLU model management day-to-day — a technical team, a CX/operations team, or a vendor-managed model? What is your expected change frequency for NLU updates (e.g., seasonal, event-driven)?"
        ),
        (
            "5.09 / 5.10 — Knowledge Base Integration",
            "Knowledge AI can ingest from many sources, but the implementation and accuracy differ significantly depending on the source system.",
            "What is your current knowledge management system (e.g., Salesforce Knowledge, ServiceNow, Confluence, SharePoint, or a custom CMS)? Is the knowledge base structured (articles with metadata) or unstructured (PDFs, web pages)? How large is the knowledge base — approximate number of articles or documents?"
        ),
        (
            "5.10 — 'Agent Opportunities' from Live Conversations",
            "The requirement to 'identify agent opportunities' from live conversations is ambiguous. It could mean coaching flags, cross-sell prompts, escalation signals, or knowledge gaps — each requires a different configuration.",
            "When you say 'identifying agent opportunities' from live conversations, what specific outcomes are you targeting — agent coaching, real-time cross-sell/upsell prompts, knowledge gap identification, or something else?"
        ),
        (
            "6.01 — Code-Switching (Language)",
            "Code-switching — where a customer mixes two languages in one utterance (e.g., English/French) — is a distinct and significantly more complex capability than supporting two separate language modes.",
            "Do you expect the system to handle mid-sentence language mixing (code-switching), or do you expect customers to interact in one language per session? Is this primarily for the Canadian French market?"
        ),
        (
            "6.04 — Language Scope",
            "The supporting criteria lists 12+ languages. Implementing NLU, TTS, STT, and full UX for 12 languages is a significant scoping and cost driver.",
            "Which of the listed languages are required at launch vs. planned for future phases? Which languages need full bot self-service, and which need only agent-side translation or transcription?"
        ),
    ],
    "Sheet 3 — Data Assessment": [
        (
            "1.04 — Third-Party System Context Capture",
            "This question asks us to preserve context from third-party systems, but the answer is entirely dependent on which systems and what data elements Ralph Lauren wants captured.",
            "Which specific third-party systems do you need context captured from during an interaction (e.g., OMS for order status, loyalty platform for VIC tier, ERP for account data)? What are the specific data fields — not just system names?"
        ),
        (
            "1.06 — Video and Fax Recording",
            "The question includes video and fax in the recording scope. These are rare in luxury retail contact centres and carry very different technical requirements.",
            "Do you currently use video interactions or fax in your contact centre operations? If so, for which use cases — advisor video calls, boutique-to-HQ, or other?"
        ),
        (
            "3.01 — Reporting Ownership",
            "The question spans ACD reporting, WFM reporting, and campaign analytics — which sit across multiple NICE products. We need to know what reporting environment Ralph Lauren expects to own and who the audience is.",
            "Who is the primary reporting audience — contact centre operations, workforce management, or business leadership? Do you have an existing BI or analytics platform (Tableau, Power BI, Looker) that reporting should feed into?"
        ),
        (
            "5.07 / 5.10 — Phone Validation & SMS Provider",
            "Phone number validation and outbound SMS delivery depend heavily on the incumbent SMS provider and whether Ralph Lauren is bringing their own short/long code or using ours.",
            "Do you have an existing SMS provider or short code? If so, which provider (Twilio, Sinch, bandwidth, etc.)? Will you bring your own SMS number or use a platform-provisioned one?"
        ),
    ],
    "Sheet 4 — Functional Assessment": [
        (
            "1.02 — Custom Context Elements",
            "The RFP mentions custom Ralph Lauren context for context preservation but does not define what that context includes. This drives integration design significantly.",
            "What specific customer attributes do you need preserved and carried across channels — e.g., VIC tier, preferred store, current cart contents, authentication status, language preference, recent order ID? Please provide a list of the fields your advisors most need at the start of an interaction."
        ),
        (
            "1.09 / 4.02 — Backend System Integrations",
            "The integrations we describe are fundamentally different depending on which OMS, CRM, ERP, and commerce platform Ralph Lauren uses. Without knowing the specific systems, we can only describe integration patterns, not actual connectors.",
            "What are the specific backend systems in scope for integration at go-live: OMS, ecommerce platform, POS/store system, loyalty/CDP, ERP, and marketing automation? Please include vendor names and versions where known."
        ),
        (
            "6.01 — Customer MFA",
            "MFA for agents is straightforward. MFA for customers during a contact centre interaction is a distinct capability with significant UX and regulatory implications. The requirement criteria mentions voice biometrics, which is a separate (Partial) capability.",
            "Is MFA for customers required at go-live, or is agent MFA the primary requirement? For customers, which authentication factors are in scope — SMS OTP, email, loyalty PIN, voice biometric, or a combination? Are there specific channels where customer MFA is not required?"
        ),
        (
            "6.03 — Customer Authentication Data",
            "The RFP asks for 'frictionless authentication for returning customers' but doesn't specify what data points Ralph Lauren uses to identify a returning customer.",
            "What identifiers do you use to recognise and authenticate returning customers today — phone number, email, loyalty ID, order number, date of birth? Which of these are you willing to use as authentication factors in IVR and chat?"
        ),
        (
            "10.01 — Customer 360 / CRM Platform",
            "The unified customer 360 view described across questions 10.01–10.04 requires a CRM integration as the system of record. The specific CRM determines what is possible natively vs. custom-built.",
            "What is your CRM of record for customer service — Salesforce Service Cloud, SAP, Microsoft Dynamics, or a custom platform? Is the Customer 360 view expected to live inside the CRM, inside the NICE agent desktop, or as a separate screen?"
        ),
        (
            "10.01 — Identity Resolution Sources",
            "Identity resolution across ecommerce, store, marketplace, and marketing requires understanding the data model and how customer identities are currently linked across channels.",
            "Is there an existing CDP or identity resolution layer (e.g., Salesforce Data Cloud, Adobe Real-Time CDP, Segment) that maintains a unified customer ID? Or does identity matching need to happen within the contact centre platform?"
        ),
        (
            "11.02 — OMS & Order Data",
            "The ability to surface and act on orders, returns, and exchanges requires a real-time OMS integration. The answer changes substantially depending on the OMS platform and whether real-time APIs are available.",
            "What is your OMS (e.g., Manhattan, SAP, Oracle, Newstore, a custom platform)? Are real-time order status APIs available, or is data batch-synced? Does the OMS have write-back APIs for return initiation and order amendments?"
        ),
        (
            "11.04 — Return & Exception Policy Rules",
            "Guiding advisors through policy exceptions requires us to either embed rules in flows or integrate with a policy/rules engine. We cannot describe this accurately without knowing where rules live.",
            "Where do your return eligibility rules and exception logic live today — in the CRM, a policy document, an internal tool, or advisor discretion? Are these rules machine-readable, or does logic need to be built into the contact centre platform from scratch?"
        ),
        (
            "11.05 — Store & Appointment Context",
            "Surfacing boutique relationships and appointment history requires integration with a store appointment and clienteling system.",
            "What system manages store appointments and boutique advisor relationships today (e.g., Salesforce, a custom tool, or a third-party clienteling platform)? Is there a real-time API to query appointment and store visit history?"
        ),
        (
            "11.07 — Guided Selling Scope",
            "Guided selling in a contact centre can mean very different things — a scripted flow for bot self-service, an agent-facing recommendation panel, or an order-on-behalf workflow. Each is a distinct capability.",
            "Is guided selling intended for the virtual agent (bot self-service), the advisor desktop (agent-assisted), or both? Do you have an existing product recommendation engine (e.g., Salesforce Commerce Cloud Einstein, Dynamic Yield, or similar) that should feed recommendations?"
        ),
        (
            "11.08 — Email Resend / Marketing Platform",
            "Resending marketing or transactional emails from the agent desktop requires an integration with the ESP or marketing automation platform.",
            "What is your email service provider or marketing automation platform (e.g., Salesforce Marketing Cloud, Braze, Klaviyo, Adobe Campaign)? Does that platform expose APIs for triggering resends or viewing email history per customer?"
        ),
        (
            "11.09 — Payment Processor",
            "PCI-safe payment completion in the contact centre requires integration with the specific payment processor or tokenisation provider. The options available differ by processor.",
            "What payment processor or gateway does Ralph Lauren use (e.g., CyberSource, Stripe, Adyen, PayPal)? Is a pay-by-link or IVR DTMF payment flow preferred, or do you need a more embedded advisor-assisted payment experience?"
        ),
        (
            "11.15 — Service Recovery Thresholds",
            "Appeasement recommendations (store credit, refund, goodwill gesture) need business rules to be meaningful. Without knowing Ralph Lauren's thresholds, we can only describe the framework, not how it would actually work.",
            "Do you have defined appeasement thresholds or recovery tiers today (e.g., order value, VIC status, number of prior complaints)? Who has approval authority for exceptions above the standard threshold?"
        ),
        (
            "11.16 — Fraud Detection System",
            "Surfacing fraud indicators in the agent desktop requires integration with your fraud detection or trust and safety platform.",
            "What fraud detection or risk system is currently in use (e.g., Signifyd, Kount, Forter, or an in-house tool)? Are fraud risk scores or flags available via API at the customer or order level?"
        ),
        (
            "11.17 — Consent Management Platform",
            "The criteria mentions OneTrust and Zeta as examples. Downstream coordination for delete and anonymisation requests depends on which platform is used.",
            "Which consent and preference management platform do you use (e.g., OneTrust, Zeta, TrustArc, or a custom solution)? Does it expose APIs for reading and writing consent status per customer identifier?"
        ),
        (
            "12.01 — Website FAQ / CMS",
            "Powering the ralphlauren.com FAQ through Knowledge AI requires understanding what CMS or content platform manages the site and how FAQ content is currently authored.",
            "What CMS or platform powers ralphlauren.com and manages FAQ/help content today? Is the FAQ currently structured (question/answer pairs with metadata) or unstructured? Who would own ongoing FAQ content authoring after go-live — a marketing team, CX ops, or a shared function?"
        ),
    ],
}

# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# Title
t = doc.add_heading("Clarification Questions for Ralph Lauren", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.font.color.rgb = NICE_NAVY
    run.font.size = Pt(18)

sub = doc.add_paragraph("NiCE — Questions Needed to Complete Accurate RFP Responses (Sheets 2, 3 & 4)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = NICE_TEAL
    run.font.size = Pt(10)
    run.font.bold = True

intro = doc.add_paragraph()
intro.paragraph_format.space_before = Pt(10)
intro.paragraph_format.space_after = Pt(6)
run = intro.add_run(
    "The following questions arose from reviewing Ralph Lauren's RFP requirements across the AI & Virtual Agent, "
    "Data, and Functional Assessment sheets. In each case, our current response is either partially accurate, "
    "placeholder-level, or would change materially based on Ralph Lauren's answer. We are requesting these "
    "clarifications before finalising our submission."
)
run.font.size = Pt(10)
run.font.color.rgb = GREY_TEXT

n = 1
for section_title, questions in QUESTIONS.items():
    # Section heading
    h = doc.add_heading(section_title, level=1)
    for run in h.runs:
        run.font.color.rgb = NICE_NAVY
        run.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after  = Pt(4)

    for (rfp_ref, why, question) in questions:
        # Question number + RFP ref
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(10)
        label.paragraph_format.space_after  = Pt(2)
        num_run = label.add_run(f"Q{n}  ")
        num_run.font.bold = True
        num_run.font.color.rgb = NICE_TEAL
        num_run.font.size = Pt(9)
        ref_run = label.add_run(f"[{rfp_ref}]")
        ref_run.font.bold = True
        ref_run.font.color.rgb = NICE_NAVY
        ref_run.font.size = Pt(10)

        # Why we need it (context box)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        set_cell_bg(cell, LIGHT_BG)
        cp = cell.paragraphs[0]
        label_run = cp.add_run("Why we need this:  ")
        label_run.font.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = NICE_NAVY
        why_run = cp.add_run(why)
        why_run.font.size = Pt(9)
        why_run.font.color.rgb = GREY_TEXT
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)

        # The actual question
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(4)
        qp.paragraph_format.space_after  = Pt(2)
        qp.paragraph_format.left_indent  = Inches(0.1)
        ql = qp.add_run("Question:  ")
        ql.font.bold = True
        ql.font.size = Pt(10)
        ql.font.color.rgb = NICE_NAVY
        qt = qp.add_run(question)
        qt.font.size = Pt(10)
        qt.font.color.rgb = GREY_TEXT

        # Answer space
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(3)
        ap.paragraph_format.space_after  = Pt(2)
        ap.paragraph_format.left_indent  = Inches(0.1)
        al = ap.add_run("Ralph Lauren response:  ")
        al.font.bold = True
        al.font.size = Pt(9)
        al.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        at = ap.add_run("_______________________________________________")
        at.font.size = Pt(9)
        at.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        n += 1

# Footer
doc.add_paragraph()
footer = doc.add_paragraph(f"NiCE Confidential  |  Ralph Lauren CCaaS RFP — Clarification Request  |  {n-1} questions across 3 sheets")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

doc.save(OUTPUT)
print(f"✓ Saved: {OUTPUT}")
print(f"  {n-1} questions across {len(QUESTIONS)} sections")
