from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- Styles ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

# --- Title block ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('Hearst Technology Services: Cognigy.AI Evaluation')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
run = subtitle.add_run('Response to Follow-Up Questions | April 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
run.italic = True

doc.add_paragraph()

# Helper: add a section heading
def add_heading(text, number):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{number}. {text}')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def add_verdict(label, color_rgb):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color_rgb
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x80, 0x50, 0x00)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x80, 0x50, 0x00)
    run2.italic = True

GREEN = RGBColor(0x1A, 0x7A, 0x3C)
AMBER = RGBColor(0xBF, 0x7A, 0x00)

# -----------------------------------------------------------------
# Q1: Communications / Campaigns
# -----------------------------------------------------------------
add_heading('Communications / Campaigns', 1)
add_verdict('Yes: via Cognigy event-triggered flows and open integration', GREEN)

add_body(
    'The observation that Cognigy agents must be explicitly activated refers to the default inbound '
    'session model; a user initiates the conversation. However, Cognigy.AI fully supports proactive '
    'and outbound communication patterns through its REST Inject API and webhook-triggered flows, '
    'which allow external systems to push a message to a user on any connected channel without the '
    'user initiating the session.'
)

add_subheading('How this works for Hearst\'s onboarding and announcement use cases:')
add_bullet('A new employee record created in Workday (or any HRIS) can trigger a Cognigy flow via webhook, sending a welcome message and onboarding checklist to the employee\'s Slack, Teams, or email channel')
add_bullet('Software migration or IT announcement campaigns can be triggered from a scheduled job or event in any upstream system (ServiceNow, Workday, a simple script) that calls the Cognigy Inject API with a contact list and message payload')
add_bullet('Follow-up messages over the first few weeks are handled by Cognigy flows with time-delayed step logic; the agent tracks where each employee is in their onboarding sequence and sends the next message at the right time')
add_bullet('Any response from the employee re-engages the full conversational AI, so a proactive message can seamlessly become a two-way support interaction')

add_body(
    'For the delivery layer (email, SMS, push), Cognigy is agnostic; it connects to whichever '
    'communication tools Hearst already uses (SendGrid, Twilio, Slack API, Teams webhooks) via '
    'its built-in extensions and HTTP request nodes. There is no requirement to replace existing '
    'notification infrastructure.'
)

# -----------------------------------------------------------------
# Q2: CSAT
# -----------------------------------------------------------------
add_heading('CSAT', 2)
add_verdict('Yes: natively within Cognigy, with open integration to any survey tool', GREEN)

add_body(
    'Cognigy.AI can collect CSAT feedback directly within the conversation without requiring a '
    'separate survey platform:'
)
add_bullet('Inline conversational feedback: at the end of any session, a Cognigy flow presents a quick rating prompt (e.g., thumbs up/down, 1-5 star rating, or NPS question) as part of the natural conversation. The response is captured as a context variable')
add_bullet('Structured survey flows: multi-question CSAT, NPS, or CES surveys can be built as a Cognigy flow branch triggered on session close, delivered on any channel (webchat, Slack, Teams, voice, SMS)')
add_bullet('Results stored and reportable: feedback scores are written to Cognigy Insights and can be forwarded to any analytics or BI tool via webhook or HTTP node for trending and correlation with resolution outcomes')

add_body(
    'For teams that already use a dedicated survey platform (e.g., Qualtrics, SurveyMonkey, Microsoft '
    'Forms), Cognigy can trigger that platform\'s survey link or API at session close. Cognigy is '
    'agnostic to the survey tool; it acts as the delivery mechanism and captures or forwards the '
    'response as needed.'
)

# -----------------------------------------------------------------
# Q3: External Knowledge
# -----------------------------------------------------------------
add_heading('External Knowledge', 3)
add_verdict('Yes: Knowledge AI supports external URL ingestion and live document crawling', GREEN)

add_body(
    'Cognigy Knowledge AI uses retrieval-augmented generation (RAG) to connect agents to external '
    'documentation. Knowledge sources can include:'
)
add_bullet('Web URLs and documentation sites: Knowledge AI crawls and ingests externally hosted documentation (e.g., Adobe\'s evergreen support docs, Microsoft support pages). The crawler handles content refresh on a configurable schedule so content stays current as vendors update it')
add_bullet('PDFs and structured documents: uploaded directly or fetched from SharePoint, Confluence, or other content repositories')
add_bullet('ServiceNow Knowledge Base: available via the certified ServiceNow integration; articles are ingested and searchable from within the bot session')

add_body(
    'At runtime, Knowledge AI retrieves the most relevant passages and grounds the LLM response in '
    'verified source content. Citations are returned with responses so users can follow the source link.'
)

add_subheading('Multi-source fallback chain')
add_body(
    'Knowledge AI supports a prioritized fallback configuration: query the Hearst ServiceNow KB first; '
    'if no match above threshold, expand to SharePoint; if still no match, fall back to a curated '
    'external source (e.g., Adobe docs, Microsoft docs). This backstop behavior directly addresses the '
    '"knowledge failover" gap identified in earlier discussions.'
)

add_note('SharePoint permissioning using delegated OAuth (user impersonation) is supported but requires configuration during deployment. This should be covered in the technical deep-dive session.')

# -----------------------------------------------------------------
# Q4: Triage
# -----------------------------------------------------------------
add_heading('Triage / ServiceNow Autonomous Action', 4)
add_verdict('Yes: Cognigy Agentic AI supports autonomous monitoring, resolution, and escalation', GREEN)

add_body(
    'Cognigy AI Agent (agentic AI) goes beyond scripted flows; it executes multi-step ServiceNow '
    'operations autonomously using tool-calling. For the ticket queue monitoring use case:'
)
add_bullet('Queue monitoring and autonomous action: The AI Agent can query open ticket queues, evaluate each ticket against a resolution playbook, and attempt automated resolution (password reset, account unlock, software provisioning) before flagging tickets it cannot resolve')
add_bullet('Classification-based escalation: When the agent cannot resolve a ticket, it classifies the issue based on intent, category, and historical assignment patterns and routes to the correct assignment group in ServiceNow, pre-populating an incident summary with no human triage required')
add_bullet('Historical classification learning: Cognigy Insights captures resolution outcomes and assignment group performance over time, feeding back into routing model refinement')

add_subheading('Example workflow for Hearst')
add_body(
    'Agent surfaces a new "printer offline" ticket, searches the KB for resolution steps, attempts '
    'an automated fix, and if unresolved, classifies it as "End User Computing > Hardware" and routes '
    'to the EUC assignment group with a pre-populated incident summary. The majority of Tier 1 tickets '
    'are resolved or routed without human triage.'
)

# -----------------------------------------------------------------
# Q5: Escalation to Live Agent
# -----------------------------------------------------------------
add_heading('Escalation to Live Agent', 5)
add_verdict('Yes: Cognigy handles escalation natively; routing layer is contact center agnostic', GREEN)

add_subheading('How escalation works')
add_body(
    'When the AI Agent determines escalation is needed (based on low confidence, an explicit user '
    'request, a failed resolution attempt, or a policy rule) it invokes Cognigy\'s built-in handover '
    'protocol. The full conversation transcript, detected intent, collected context variables (employee '
    'ID, ticket number, issue category), and sentiment signals are bundled into a handover payload and '
    'passed to the receiving live agent system.'
)

add_body(
    'Cognigy ships with out-of-the-box handover integrations for the most common agent desktop and '
    'contact center platforms. For any platform not on the built-in list, Cognigy\'s open handover '
    'API makes it straightforward to connect; Cognigy is agnostic to the contact center or '
    'service desk routing layer Hearst uses today or in the future.'
)

add_subheading('Visibility and reporting')
add_bullet('Escalation events are captured in Cognigy Insights, filterable by channel, intent, escalation reason, and time period')
add_bullet('Containment rate (bot-resolved vs. escalated) is a first-class metric in Cognigy Insights, enabling managers to see exactly where automation succeeds and where it falls short')
add_bullet('The full transcript is available for every escalated session for QA review')
add_bullet('Escalation rates can be exported via API into any BI or analytics tool for trending alongside ServiceNow resolution data')

# -----------------------------------------------------------------
# Q6: Analytics
# -----------------------------------------------------------------
add_heading('Analytics', 6)
add_verdict('Yes: Cognigy Insights is the native analytics layer, with open export to any BI tool', GREEN)

add_subheading('What Cognigy Insights provides today')
add_bullet('Session volumes, channel breakdown, containment rate, and escalation rate across all connected channels (webchat, Slack, Teams, voice, SMS)')
add_bullet('Flow step analytics: see where users drop off, which flow branches are taken most, and where the agent falls back to clarification')
add_bullet('Intent distribution and NLU performance: which intents are firing, which are misclassified, and where training gaps exist')
add_bullet('Utterance-level data: raw user inputs are captured and queryable, supporting sentiment trend analysis beyond simple thumbs up/down')
add_bullet('Historical data accessible via the Insights UI and via Cognigy\'s OData API; data can be pulled into Tableau, Power BI, Snowflake, or any BI tool for multi-year trending and custom dashboards')

add_subheading('What was discussed in the earlier review session')
add_body(
    'The concern raised about analytics being "replaced" or in transition is understood. Cognigy\'s '
    'Insights product is the stable, production analytics layer for the conversational AI layer. '
    'For any broader contact center or workforce analytics needs, Cognigy is open; it exports all '
    'interaction data via API and integrates with whichever analytics platform Hearst standardizes on. '
    'There is no lock-in to a specific analytics tool.'
)
add_note('Confirm with your Cognigy account team: data retention window for Insights, and whether the OData export covers the full historical range Hearst requires (3+ years).')

# -----------------------------------------------------------------
# Q7: Autonomous Use Case Identification
# -----------------------------------------------------------------
add_heading('Autonomous Use Case Identification', 7)
add_verdict('Partial: analytics-driven discovery is available; automated backlog generation requires a human review step', AMBER)

add_body(
    'Cognigy provides strong building blocks for use case identification, though fully automated '
    'backlog generation (as Moveworks positions its use case pipeline) is not a single '
    'out-of-the-box feature.'
)

add_subheading('What is available')
add_bullet('Cognigy Intent Trainer reviews sessions where the AI fell back or misclassified, categorizes patterns of unhandled utterances, and surfaces them for review; these are direct signals for new automation candidates')
add_bullet('Containment gap analysis: the delta between bot-handled and escalated interactions, broken down by intent and topic, is a natural use case backlog surfaced in Cognigy Insights and exportable via API')
add_bullet('Cognigy Insights surfaces dead-end flow paths and unrecognized input patterns, providing the raw material for identifying what to build next')
add_bullet('Topic clustering from conversation transcripts: recurring themes in agent-handled or escalated sessions highlight where new automation would have the most impact')

add_subheading('What requires a human step')
add_body(
    'Translating those signals into a prioritized backlog with effort estimates and solution '
    'recommendations is not fully automated; a human (conversation designer or process owner) '
    'reviews the surfaced patterns and decides what to build next. Cognigy\'s Professional Services '
    'team can facilitate this as a periodic automation opportunity review engagement.'
)

add_note('Recommended cadence for Hearst: quarterly review using Cognigy Insights and Intent Trainer outputs to identify the top 5 new automation candidates per cycle.')

# -----------------------------------------------------------------
# Summary table
# -----------------------------------------------------------------
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run('Summary')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
p.paragraph_format.space_after = Pt(8)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'

# Header row
hdr = table.rows[0]
for cell in hdr.cells:
    cell._tc.get_or_add_tcPr()
for i, text in enumerate(['Question', 'Capability', 'Key Note']):
    cell = hdr.cells[i]
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    tcPr.append(shd)

rows_data = [
    ('Outbound Campaigns', 'Yes: Cognigy Inject API + open integration', 'Cognigy triggers flows; delivery layer connects to existing tools (Slack, Twilio, etc.)'),
    ('CSAT Surveys', 'Yes: native in Cognigy + open integration', 'Inline conversational CSAT or trigger to existing survey platform'),
    ('External Knowledge', 'Yes: Knowledge AI RAG', 'URL crawling, multi-source fallback, SharePoint permissioning configurable'),
    ('ServiceNow Triage', 'Yes: Cognigy Agentic AI', 'Autonomous resolution + classification-based escalation'),
    ('Escalation Visibility', 'Yes: Cognigy handover + open integration', 'Full transcript handoff; contact center routing layer is platform agnostic'),
    ('Analytics', 'Yes: Cognigy Insights + open export', 'Insights is production; OData API feeds any BI tool for multi-year trending'),
    ('Use Case Identification', 'Partial', 'Intent Trainer + Insights surfaces candidates; human review step for backlog'),
]

for i, (q, cap, note) in enumerate(rows_data):
    row = table.rows[i + 1]
    row.cells[0].text = q
    row.cells[1].text = cap
    row.cells[2].text = note
    if i % 2 == 0:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EEF3F9')
            tcPr.append(shd)

widths = [Inches(1.8), Inches(2.2), Inches(3.0)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Follow-up actions: ')
run.bold = True
run2 = p.add_run(
    'Two items warrant follow-up with the Cognigy account team before committing to scope: '
    '(1) Insights data retention window and OData export range for the 3+ year history requirement; '
    '(2) SharePoint permissioning via delegated OAuth (user impersonation) is supported but should be '
    'confirmed and covered in the technical deep-dive session.'
)
run2.italic = True

output_path = '/Users/Adam.Boyle/Cognigy/rfp/Hearst-RFP-Responses-April2026.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
