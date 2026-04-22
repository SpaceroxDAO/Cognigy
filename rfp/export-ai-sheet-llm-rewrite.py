"""
AI & Virtual Agent Assessment — Rewrite for Ralph Lauren RFP
AI-based understanding is treated as the natural default throughout.
No em dashes. No "LLM-first" framing. Demo scenario references included.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DOCX = "/Users/Adam.Boyle/Downloads/NiCE - Ralph Lauren AI Assessment (LLM-First).docx"

NICE_NAVY = RGBColor(0x00, 0x29, 0x5E)
NICE_TEAL = RGBColor(0x00, 0x97, 0x9D)
AMBER     = RGBColor(0xC4, 0x73, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT = RGBColor(0x22, 0x22, 0x22)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Q&A data ──────────────────────────────────────────────────────────────────
# (id, subsection, capability, question, compliance, response, citation)
QA = [

    # ── 1. Model Governance ───────────────────────────────────────────────────

    ("1.01", "Model Governance", "Model Transparency",
     "Ability to provide transparent documentation of AI/ML models, including decision logic and model architecture",
     "Yes",
     "NICE Cognigy's AI agent operates from a fully readable system prompt that defines the agent's persona, rules, and behaviour logic. This is the model's decision logic in practice: every instruction is human-readable, version-controlled, and auditable. The AI provider is explicitly configured and documented at the tenant level. Prompt version history, deployment audit trails, and configuration change logs are all maintained in the admin console and exportable for compliance review. In the Admin Configuration demo scenario (2.5), this is visible through the virtual agent configuration and sandbox promotion workflow.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("1.02", "Model Governance", "Model Transparency",
     "Ability to avoid 'black box' models by providing explainability features for all AI-driven decisions, supporting comprehensive audit trails",
     "Yes",
     "NICE Cognigy makes conversational logic transparent by design. The system prompt is fully human-readable and auditable, with no hidden classifiers or opaque decision trees. Every interaction generates a conversation trace showing the AI's inputs, tool calls invoked, and outputs returned. Ralph Lauren's team can inspect, modify, and version-control the reasoning logic directly through the admin console. When Knowledge AI is used, each response includes citations showing which source passages grounded the answer.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("1.03", "Model Governance", "Model Training, Validation, and Monitoring",
     "Ability to assess and document model training and validation processes, including evaluation of model performance, bias, and accuracy",
     "Yes",
     "The proposed deployment for Ralph Lauren uses a pre-trained AI foundation model selected by Ralph Lauren, with no custom NLU training cycle required. Validation is performed through NICE Cognigy's Playbook testing framework, where scripted test scenarios are run against the live agent configuration and scored for accuracy. Ongoing performance is monitored through NICE CXone conversation analytics: escalation rates, resolution rates, and topic distribution are tracked continuously to identify prompt improvements or knowledge base gaps. The AI provider also publishes model cards and evaluation benchmarks for each model release.",
     "https://docs.cognigy.com/ai/test/playbooks/"),

    ("1.04", "Model Governance", "Bias, Fairness, and Ethical Considerations",
     "Ability to identify, mitigate, and monitor bias in AI/ML models, and demonstrate adherence to ethical standards and fairness",
     "Partial",
     "Bias mitigation in the proposed deployment operates at three levels. At the AI provider level, all supported providers (OpenAI, Anthropic, Google, and others) publish Responsible AI policies covering bias evaluation, content filtering, and fairness standards; this is the provider's responsibility. At the retrieval level, NICE Cognigy Knowledge AI constrains AI responses to verified, curated content, significantly reducing the risk of biased or ungrounded generation. At the application level, system prompt guardrails, tool call restrictions, and escalation thresholds limit the scope of autonomous AI decisions. NICE Cognigy does not provide its own bias testing suite for foundation models; that responsibility sits with the selected AI provider.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    # ── 2. IVR Self-Service ───────────────────────────────────────────────────

    ("2.01", "IVR-Self Service", "NLU/NLP (Cognitive Engagement)",
     "Ability to drive conversational interactions with customers through automation (virtual assistant/self-service) without the need for human CSRs",
     "Yes",
     "NICE Cognigy drives conversational self-service using a large language model as the core understanding and reasoning engine. The AI processes customer utterances directly, understands free-form requests, extracts relevant data, and decides which tools to invoke, with no intent training or example utterances required. In the Where Is My Order scenario (1.01), the virtual agent authenticates the caller by name, performs a real-time order lookup via IBM Sterling OMS, surfaces the latest shipment status, and presents self-service resolution options, all without agent involvement.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("2.02", "IVR-Self Service", "Proactive Notifications",
     "Ability to proactively initiate outbound IVR or voice bot calls to customers for task completion (e.g., confirming a delivery, appointment reminders)",
     "Yes",
     "NICE CXone SmartReach, acquired via LiveVox in 2024, delivers proactive omnichannel outbound campaigns across voice, SMS, email, and WhatsApp. Event-triggered campaigns fire from CRM or OMS events, covering abandoned cart follow-up, delivery updates, return authorisations, appointment reminders, and promotional awareness. In the BOPIS scenario (1.02), SmartReach sends a ready-for-pickup notification when the store confirms item availability, and a post-purchase outreach message with styling suggestions. Consent management, DNC scrubbing, and TCPA/CTIA compliance automation are all built in.",
     "https://www.nice.com/products/cxone-smartreach"),

    ("2.03", "IVR-Self Service", "Outbound Engagement",
     "Ability to route outbound interactions to a contact center representative when required",
     "Yes",
     "Outbound interactions requiring live agent handling are routed to NICE CXone's ACD with full context preservation. The SmartReach campaign engine blends outbound agent connections into the universal queue alongside inbound interactions, with real-time pacing controls for supervisor oversight.",
     "https://www.nice.com/products/cxone-smartreach"),

    # ── 3. Voice / Chat Bots ─────────────────────────────────────────────────

    ("3.01", "Voice / Chat bots", "Voice / Chat bots",
     "Ability to provide chatbot services to provide intelligent assistance capabilities",
     "Yes",
     "NICE Cognigy deploys AI-powered chatbots across web, mobile, and messaging channels. NICE Cognigy Webchat v3 is WCAG 2.1 AA accessible. In the Returns scenario (1.03) and Guided Selling scenario (1.02), the virtual agent handles structured data collection, eligibility checking, policy lookups, and product recommendations entirely in self-service, escalating to a live CSR only when required and preserving full context at the point of handoff.",
     "https://docs.cognigy.com/webchat/v3/overview"),

    ("3.02", "Voice / Chat bots", "Voice / Chat bots",
     "Ability to build and deploy natural language voicebots to replace self-service (IVR) menus",
     "Yes",
     "Natural language voicebots that replace IVR menus are built and deployed through NICE Cognigy's Voice Gateway and visual flow builder. The AI handles all conversational understanding, so customers speak naturally without being told to use specific keywords or navigate menu trees. In the Where Is My Order scenario (1.01), the caller simply says 'Where is my order?' and the virtual agent authenticates, performs the OMS lookup, and presents resolution options in a natural conversation with no DTMF menus.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("3.03", "Voice / Chat bots", "NLU/NLP (Cognitive Engagement)",
     "Ability to use NLU and AI-powered solutions to partially or fully automate reasonably complex responses to customer inquiries",
     "Yes",
     "Complex customer inquiries are automated through NICE Cognigy's AI agent. The AI processes free-form input, determines the appropriate action, invokes the relevant tools (order lookup, account verification, return initiation, Knowledge AI retrieval), and generates natural, grounded responses. In the Returns scenario (1.03), the AI checks return eligibility against policy rules, validates via IBM Sterling OMS, performs a fraud check, and presents resolution options, all in self-service with no agent involvement.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("3.04", "Voice / Chat bots", "NLU/NLP (Cognitive Engagement)",
     "Ability to integrate with third-party NLU engines (Dialogflow/CCAI, Azure, AWS Lex, Decagon, IBM Watson, Kore.ai, etc.)",
     "Yes",
     "The proposed architecture for Ralph Lauren uses a pre-trained AI model for all natural language understanding, without a separate NLU intent classifier. NICE Cognigy's NLU Connector framework supports Google Dialogflow, Amazon Lex, IBM Watson, Kore.ai, and custom endpoints for hybrid configurations if required. The capability is fully available; it simply is not part of the proposed architecture.",
     "https://docs.cognigy.com/ai/empower/nlu/nlu-connectors/overview/"),

    ("3.05", "Voice / Chat bots", "NLU/NLP (Cognitive Engagement)",
     "Ability to support directed dialogue or guided speech (e.g., customers can use specific words to indicate a choice or selection)",
     "Yes",
     "NICE Cognigy's AI agent handles fully free-form speech: customers do not need to follow a guided prompt structure or use specific words. The AI understands natural phrasing and contextual intent from any utterance. Where structured data collection is required, such as capturing an order number in the Returns scenario or a size preference in the Guided Selling scenario, the agent asks naturally and the AI parses the response. Traditional directed dialogue menus remain configurable for fallback flows or compliance-driven scenarios.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("3.06", "Voice / Chat bots", "Rep Assist / Agent Assist",
     "Ability to deploy virtual assistant(s) in support of CSRs to expedite customer call / chat resolution",
     "Yes",
     "NICE Cognigy Agent Copilot provides real-time AI assistance during live interactions, including next-best-action suggestions, Knowledge AI content surfaced automatically, and AI-drafted response suggestions for agents to review and send. In the Where Is My Order scenario (1.01), Agent Copilot detects customer frustration from the live transcript and surfaces a promotional discount as a service recovery recommendation, which the CSR applies in one click. In the Guided Selling scenario (1.02), Agent Copilot surfaces active promotions and recommends complementary items based on the customer's stated preferences.",
     "https://www.cognigy.com/platform/agent-copilot"),

    ("3.07", "Voice / Chat bots", "Proactive Notifications",
     "Ability to trigger notifications/windows/actions for customers based on monitored events",
     "Yes",
     "Event-driven customer notifications, covering order status changes, delivery exceptions, loyalty milestones, and OMS-triggered service recovery, are orchestrated through NICE Cognigy's integration with CRM and OMS event streams. In the Where Is My Order scenario (1.01), the platform sends an SMS or email confirmation summarising the shipment update and the applied promotion. In the BOPIS scenario (1.02), a ready-for-pickup notification fires automatically when the store confirms availability. Outbound voice notifications are delivered via NICE CXone SmartReach.",
     "https://www.nice.com/products/cxone-smartreach"),

    ("3.08", "Voice / Chat bots", "Speech/Text Analytics",
     "Ability to capture speech analytics including voice transcription and interaction overview to organize unstructured data",
     "Yes",
     "NICE CXone Interaction Analytics transcribes 100% of voice interactions using integrated STT providers and runs AI-powered topic detection, sentiment scoring, and category tagging across all recordings. Unstructured voice data is automatically organised into searchable, structured interaction records with speaker attribution and timestamped segments. In the back-office panel scenario (2.3 and 2.4), supervisors access these analytics through the QM workflow and reporting dashboards.",
     "https://www.nice.com/products/interaction-analytics"),

    ("3.09", "Voice / Chat bots", "Speech/Text Analytics",
     "Ability to capture text analytics including voice transcription, sentiment analysis, and interaction overview to organize unstructured data",
     "Yes",
     "NICE CXone Interaction Analytics applies AI-powered sentiment analysis and category tagging across chat, email, and transcribed voice interactions. AI-generated interaction summaries provide natural language overviews of each interaction, making unstructured data actionable for QA, coaching, and product teams without manual review.",
     "https://www.nice.com/products/interaction-analytics"),

    ("3.10", "Voice / Chat bots", "Speech/Text Analytics",
     "Ability to automate scoring for call, chat, and other text-based interactions based on behavioral analytics and NLP",
     "Yes",
     "NICE CXone Auto QA evaluates 100% of interactions against configurable quality scorecards using AI-powered behavioural analytics. Scores, coaching flags, and team-level trend analysis are surfaced automatically to supervisors, eliminating manual QA sampling and ensuring consistent quality measurement across all channels. In the back-office panel scenario (2.3), the QM workflow demonstrates automated scoring, evaluation rubrics, and the agent feedback loop.",
     "https://www.nice.com/products/quality-management"),

    # ── 4. STT / TTS ─────────────────────────────────────────────────────────

    ("4.01", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to support multi-language speech-to-text for speech recognition, call transcription, sentiment analysis, and analytics",
     "Yes",
     "Multi-language speech-to-text is delivered through NICE Cognigy's Voice Gateway integrations with Deepgram Nova-3 (primary), Azure Speech Services, and Google Cloud STT, all supporting multi-language transcription, sentiment-annotated transcripts, and analytics. Language selection is automatic from the first utterance or manually configured per interaction. In the Language Translation return scenario (1.03L), the Swedish-speaking customer is automatically detected and the correct STT configuration is applied without manual intervention.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.02", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to perform speaker diarization when there are more than two channels of interaction, accurately attributing speech to the correct speaker",
     "Partial",
     "Speaker diarisation for multi-channel interactions is available through Deepgram Nova-3 (primary STT provider), which supports multi-speaker diarisation natively. Azure Speech Services is available as an alternative with diarisation support. Not all supported STT providers include diarisation, so availability depends on the configured provider for the deployment.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.03", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to tune native ASR/STT with recordings or IPA definitions",
     "Yes",
     "Custom vocabulary tuning for brand names, product lines, and luxury retail terminology is supported through the STT provider's customisation capabilities. Deepgram supports custom language model fine-tuning using training audio. Azure Speech Services supports custom pronunciation lexicons using IPA definitions. Tuning is managed within the STT provider's tooling and referenced from NICE Cognigy's Voice Gateway configuration.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.04", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to automate transcription of post-call notes from STT with summarization, including the ability to customize post-call note output",
     "Yes",
     "NICE Cognigy generates post-call summaries by passing the interaction transcript to the configured AI model with a fully customisable summarisation prompt. The AI produces structured summaries matching Ralph Lauren's case management field format, covering disposition, order number, action taken, and follow-up required, and writes them back to the OMS or CRM via API at the end of each interaction. In the Where Is My Order scenario (1.01), ACW fields are pre-populated within seconds of the call ending.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("4.05", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to transcribe text-based interactions and format them as conversations between system/agent and customer",
     "Yes",
     "NICE CXone automatically transcribes text-based interactions from chat, email, and messaging channels and structures them as timestamped conversational records with full speaker attribution. Records are searchable, stored per configured retention policy, and accessible via API for downstream integration.",
     "https://www.nice.com/products/recording"),

    ("4.06", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability to transcribe calls and chat interactions and analyze call transcripts",
     "Yes",
     "Voice and text interactions are transcribed by NICE CXone Interaction Analytics with full-text search across transcripts. AI-powered summarisation and topic extraction enhance raw transcripts into structured, actionable intelligence for QA teams, supervisors, and operational reporting.",
     "https://www.nice.com/products/interaction-analytics"),

    ("4.07", "Speech to Text and Text to Speech", "Speech to Text",
     "Ability of ASR/STT models to detect and filter background noise or multiple speakers during conversations, with explicit controls",
     "Yes",
     "NICE Cognigy's Voice Gateway processes audio through Deepgram Nova-3 with configurable noise suppression, endpointing sensitivity, and multi-speaker handling. Background noise filtering parameters, including silence thresholds and audio enhancement settings, are tunable per deployment through the Voice Gateway configuration.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.08", "Speech to Text and Text to Speech", "Text to Speech",
     "Ability to provide / integrate / access text-to-speech services",
     "Yes",
     "Text-to-speech output is delivered through NICE Cognigy's Voice Gateway TTS integrations: ElevenLabs eleven_multilingual_v2 (primary, natural multilingual neural voice), Microsoft Azure Neural TTS, Google WaveNet, and Nuance. Multiple TTS providers can run simultaneously with prioritisation rules for failover or A/B persona testing.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.09", "Speech to Text and Text to Speech", "Text to Speech",
     "Ability to convert system text into speech for IVR and conversational interface outputs",
     "Yes",
     "NICE Cognigy's Voice Gateway converts all agent responses, including AI-generated text, to natural-sounding speech through the configured TTS provider. AI responses are streamed to TTS in real time for low-latency voice output. The same pipeline serves both self-service bot conversations and agent-scripted prompts.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.10", "Speech to Text and Text to Speech", "Text-to-Speech",
     "Ability to customize native or service provider-specific text-to-speech (synthetic) voices and personas, with support for multiple simultaneous TTS engines",
     "Yes",
     "Custom voice personas are configured per deployment in NICE Cognigy's Voice Gateway. ElevenLabs supports bespoke voice cloning and persona creation, enabling Ralph Lauren to establish a consistent branded voice across all self-service touchpoints. Pronunciation dictionaries handle brand names, product lines, and luxury retail vocabulary across all TTS providers. Multiple TTS engines can run simultaneously with priority ordering.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("4.11", "Speech to Text and Text to Speech", "Text Transcription",
     "Ability to generate readable, searchable, timestamped call transcripts",
     "Yes",
     "Readable, searchable, and timestamped call transcripts are generated automatically for all voice interactions processed through NICE Cognigy's Voice Gateway. Transcripts are stored in NICE CXone with full-text search, speaker labels, and sentiment annotation, and are accessible via the analytics console and retrievable via API for case management or compliance workflows.",
     "https://www.nice.com/products/interaction-analytics"),

    # ── 5. AI Model & Conversational Capabilities ─────────────────────────────

    ("5.01", "AI Model & Conversational Capabilities", "AI Model & Conversational Capabilities",
     "Ability to manage and update AI models efficiently",
     "Yes",
     "NICE Cognigy's AI management workflow covers system prompts, Knowledge AI content, tool definitions, and AI provider configuration. System prompts are version-controlled, staged in a sandbox environment, and promoted to production through role-based approval workflows. Provider updates such as model version changes are applied through the AI Connector settings and validated in sandbox before production deployment. In the Admin Configuration scenario (2.5), this workflow is demonstrated end-to-end, covering virtual agent configuration, knowledge source management, and non-production environment promotion.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("5.02", "AI Model & Conversational Capabilities", "AI Model & Conversational Capabilities",
     "Ability to understand and respond to complex, multi-turn conversations",
     "Yes",
     "Multi-turn conversations requiring complex reasoning are handled natively by NICE Cognigy's AI integration layer. Session context is maintained in the Cognigy Context object across all turns, giving the AI full access to conversation history, extracted customer data, and prior tool call results at every step. In the Guided Selling scenario (1.02), the AI manages preference gathering, product recommendation, promotion application, and BOPIS order capture across multiple turns without losing context. In the Language Translation return scenario (1.03L), context is preserved from the initial Swedish-language chat through to the English-speaking CSR.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("5.03", "AI Model & Conversational Capabilities", "AI Model & Conversational Capabilities",
     "Ability to escalate conversations to a human CSR when necessary",
     "Yes",
     "Escalation to live CSRs is managed through NICE Cognigy's built-in handover protocol, which queues the interaction to NICE CXone's ACD with the full transcript, conversation context, extracted customer data, and intent history pre-populated on the agent screen before the call connects. In the Where Is My Order scenario (1.01), the CSR receives the full context payload before greeting the customer so the customer does not need to repeat any information.",
     "https://docs.cognigy.com/ai/escalate/handovers"),

    ("5.04", "AI Model & Conversational Capabilities", "AI Model & Conversational Capabilities",
     "Ability to learn and improve from user interactions (continuous learning)",
     "Yes",
     "Continuous improvement works through Knowledge AI content updates, system prompt refinement, and retrieval index enrichment rather than intent retraining. NICE Cognigy's conversation analytics surface topics that were escalated, unanswered, or resolved poorly. These insights inform Knowledge AI document updates and prompt improvements, and as Knowledge AI is updated, the retrieval pipeline incorporates the new content immediately with no deployment cycle required. AI model performance also improves as the provider releases updated versions, adopted through the AI Connector version setting.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("5.05", "AI Model & Conversational Capabilities", "NLU/NLP (Cognitive Engagement)",
     "Ability to enable Ralph Lauren-led training, management, and governance of NLU models for AI-powered flows (voice, chat, messaging)",
     "Yes",
     "Ralph Lauren governs the AI agent's behaviour through system prompt management, Knowledge AI content governance, and AI provider configuration. Role-based access controls determine who can edit prompts and deploy changes. Sandbox-to-production promotion workflows and full audit logs provide a complete governance chain. In the Admin Configuration scenario (2.5), this is demonstrated through virtual agent configuration, knowledge source management, and the non-production test environment.",
     "https://docs.cognigy.com/ai/empower/agentic-ai/overview/"),

    ("5.06", "AI Model & Conversational Capabilities", "NLU/NLP (Cognitive Engagement)",
     "Ability to support conversational / natural language-based interactions; customers can use free-form sentences to denote requirements",
     "Yes",
     "NICE Cognigy's AI agent handles free-form customer input as the default interaction mode across all three Ralph Lauren demo scenarios: voice self-service in the order inquiry (1.01), guided chat in the sales and BOPIS flow (1.02), and digital messaging in the returns flow (1.03). Customers speak or type in any natural phrasing and the AI understands intent directly, with no keyword matching or trained classifiers required.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("5.07", "AI Model & Conversational Capabilities", "Speech to Text",
     "Ability to host or integrate with LLMs provided by third-party service providers",
     "Yes",
     "NICE Cognigy's AI Connector integrates with all leading providers, including OpenAI, Azure OpenAI, Anthropic Claude, Google Gemini, and Amazon Bedrock, plus any OpenAI-compatible endpoint. NICE Cognigy is fully model-agnostic: Ralph Lauren selects and configures the preferred provider. API keys, model version selection, and request routing remain under Ralph Lauren's control.",
     "https://docs.cognigy.com/ai/empower/llms/overview/"),

    ("5.08", "AI Model & Conversational Capabilities", "Speech to Text",
     "Ability to host fine-tuned large language models (LLMs)",
     "Yes",
     "NICE Cognigy does not host AI models. Any OpenAI-compatible fine-tuned model, whether hosted by Ralph Lauren on a private endpoint, on a cloud provider's fine-tuning service, or through a third party, can be registered in NICE Cognigy's AI Connector as a custom endpoint and used as the primary reasoning model for the deployment.",
     "https://docs.cognigy.com/ai/empower/llms/overview/"),

    ("5.09", "AI Model & Conversational Capabilities", "LLM / Gen AI",
     "Ability for the virtual agent / chatbot to search and learn from the knowledge base to provide answers for customer questions",
     "Yes",
     "NICE Cognigy Knowledge AI ingests PDFs, URLs, and structured documents, chunks and embeds them, and retrieves relevant passages at runtime to ground AI responses. The AI generates answers from retrieved passages rather than from training data, preventing hallucination and ensuring every response reflects Ralph Lauren's current policies, products, and procedures. Citations are returned with each answer. In the Returns scenario (1.03), the virtual agent uses Knowledge AI to validate return policy eligibility in real time.",
     "https://www.cognigy.com/platform/cognigy-knowledge-ai"),

    ("5.10", "AI Model & Conversational Capabilities", "LLM / Gen AI",
     "Ability to improve the knowledge base and subsequent LLM-generated content by actively listening to customer-live agent conversations",
     "Yes",
     "Conversation analytics in NICE Cognigy surfaces topics that were escalated, unanswered, or handled poorly, which are flagged for Knowledge AI content review. Ralph Lauren's team adds or updates KB documents based on these insights. As Knowledge AI is updated, the retrieval pipeline incorporates the new content immediately with no retraining or deployment cycle required.",
     "https://www.cognigy.com/platform/cognigy-knowledge-ai"),

    # ── 6. Language Support ───────────────────────────────────────────────────

    ("6.01", "Language Support and Enablement", "Language Support",
     "Ability to support the following in English and French in the native solution: NLU/NLP, TTS, STT",
     "Yes",
     "English and French are fully supported with no separate training or language models required. The AI understands and responds natively in both languages through the same model. TTS: ElevenLabs eleven_multilingual_v2 delivers natural-sounding voice output in both languages. STT: Deepgram Nova-3 transcribes both languages with high accuracy. Language detection is automatic from the first utterance. The Language Translation scenario (1.03L) demonstrates the same detection and routing mechanism applied to Swedish, confirming the extensibility of the approach across Ralph Lauren's global markets.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("6.02", "Language Support and Enablement", "Language Support",
     "Ability to automate translation of calls, chats, and other text-based interactions, including leveraging native Speech to Text for transcription and summary",
     "Yes",
     "NICE Cognigy supports automated translation of voice and chat interactions through its AI layer. Where a customer speaks in a language outside the primary configured set, the AI translates the input, responds in the customer's language, and simultaneously provides the agent with an English transcript and summary. In the Language Translation return scenario (1.03L), the CSR speaks English while the Swedish-speaking customer hears responses in Swedish, with live bidirectional translation managed automatically throughout the interaction.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("6.03", "Language Support and Enablement", "Language Support",
     "Ability to automatically detect and manually select language that is to be translated",
     "Yes",
     "Automatic language detection is built into NICE Cognigy's STT pipeline. The first utterance determines the language, which selects the appropriate STT configuration and triggers the correct AI language context. Agents and supervisors can manually override the language setting at any point during the interaction. In the Language Translation scenario (1.03L), Swedish is detected automatically and the translation layer is activated without any agent configuration step.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),

    ("6.04", "Language Support and Enablement", "Language Support",
     "Ability to provide consistent user (customer and employee) experience across all supported languages",
     "Yes",
     "NICE Cognigy provides a consistent experience across English and French because the same AI model handles both languages through the same flows, knowledge base, and tool set. There are no separate personas or logic paths per language. Pronunciation dictionaries and TTS voice configurations are maintained per language for consistent and natural audio output. In the Language Translation scenario (1.03L), the same underlying flow and knowledge base serves the Swedish-speaking customer without a separate configuration.",
     "https://www.cognigy.com/platform/cognigy-ai"),

    ("6.05", "Language Support and Enablement", "Language Support",
     "Ability to monitor and report on language usage and performance",
     "Yes",
     "Language usage and performance reporting, covering interaction volume by language, escalation rates, sentiment trends, and resolution rates, is available through NICE CXone Interaction Analytics dashboards. Reports are exportable via API for integration with Ralph Lauren's BI tools. The back-office panel scenario (2.4) demonstrates the custom report builder and standard reporting library.",
     "https://www.nice.com/products/interaction-analytics"),

    ("6.06", "Language Support and Enablement", "Language Support",
     "Ability to support language-specific regulatory and compliance requirements",
     "Yes",
     "NICE CXone's compliance framework supports language and region-specific requirements including GDPR (EU/FR), HIPAA (US), CCPA, and PIPEDA (Canada). Data residency is configurable per region, and French-language deployments can be hosted in the EU Frankfurt data centre for full GDPR isolation. Regulatory change monitoring is conducted continuously by NICE's legal and compliance team, with platform updates applied proactively.",
     "https://www.nice.com/company/legal/gdpr-faqs"),

    # ── 7. AI Integrations ────────────────────────────────────────────────────

    ("7.01", "AI Integrations", "AI Integrations",
     "Ability to integrate with third-party NLU/NLP engines",
     "Yes",
     "The proposed architecture for Ralph Lauren uses the AI model directly for all natural language understanding, without a separate NLU intent classifier. NICE Cognigy's NLU Connector framework supports Google Dialogflow, Amazon Lex, IBM Watson Assistant, Kore.ai, and custom endpoints for hybrid or benchmarking configurations if required. The capability is available and documented; it is simply not part of the proposed architecture.",
     "https://docs.cognigy.com/ai/empower/nlu/nlu-connectors/overview/"),

    ("7.02", "AI Integrations", "AI Integrations",
     "Ability to integrate with third-party TTS engines",
     "Yes",
     "Third-party TTS integration, covering ElevenLabs (primary), Azure Neural TTS, Google WaveNet, and Nuance, is managed through NICE Cognigy's Voice Gateway. Multiple TTS providers can be configured simultaneously with priority ordering for failover or A/B persona testing. Provider configuration is managed per deployment without code changes.",
     "https://www.cognigy.com/platform/cognigy-voice-gateway"),
]

# ── Build document ────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

title = doc.add_heading("AI and Virtual Agent Assessment", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = NICE_NAVY
    run.font.size = Pt(20)

sub = doc.add_paragraph("NiCE Response  |  Ralph Lauren CCaaS RFP")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = NICE_TEAL
    run.font.size = Pt(11)
    run.font.bold = True

doc.add_paragraph()

yes_count     = sum(1 for q in QA if q[4] == "Yes")
partial_count = sum(1 for q in QA if q[4] == "Partial")
stats = doc.add_paragraph()
stats.paragraph_format.space_after = Pt(16)
stats.add_run(f"{len(QA)} requirements   ").font.size = Pt(9)
yr = stats.add_run(f"Yes: {yes_count}   ")
yr.font.size = Pt(9); yr.font.bold = True; yr.font.color.rgb = NICE_NAVY
pr = stats.add_run(f"Partial: {partial_count}")
pr.font.size = Pt(9); pr.font.bold = True; pr.font.color.rgb = AMBER

current_subsec = None
for (req_id, subsec, cap, question, compliance, response, citation) in QA:
    if subsec != current_subsec:
        current_subsec = subsec
        h = doc.add_heading(subsec, level=1)
        for run in h.runs:
            run.font.color.rgb = NICE_NAVY
            run.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after  = Pt(4)

    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(9)
    label.paragraph_format.space_after  = Pt(2)
    id_run = label.add_run(f"{req_id}  ")
    id_run.font.bold = True; id_run.font.color.rgb = NICE_TEAL; id_run.font.size = Pt(8.5)
    cap_run = label.add_run(cap)
    cap_run.font.bold = True; cap_run.font.color.rgb = NICE_NAVY; cap_run.font.size = Pt(9.5)

    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(1)
    qp.paragraph_format.space_after  = Pt(4)
    qp.paragraph_format.left_indent  = Inches(0.1)
    qr = qp.add_run(question)
    qr.font.size = Pt(9.5); qr.font.color.rgb = DARK_TEXT

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.allow_autofit = False

    badge_cell = tbl.cell(0, 0)
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(compliance)
    br.font.bold = True; br.font.size = Pt(9); br.font.color.rgb = WHITE
    set_cell_bg(badge_cell, "00295E" if compliance == "Yes" else "C47300")

    resp_cell = tbl.cell(0, 1)
    rp = resp_cell.paragraphs[0]
    rr = rp.add_run(response)
    rr.font.size = Pt(9.5)

    if citation:
        cp = resp_cell.add_paragraph()
        cr = cp.add_run(f"  {citation}")
        cr.font.size = Pt(8); cr.font.color.rgb = NICE_TEAL; cr.font.italic = True

    tbl.columns[0].width = Inches(0.7)
    tbl.columns[1].width = Inches(5.4)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()
footer = doc.add_paragraph("Prepared by NICE  |  Confidential  |  Ralph Lauren CCaaS RFP  |  April 2026")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8); run.font.color.rgb = GRAY; run.font.italic = True

doc.save(OUTPUT_DOCX)
print(f"Saved: {OUTPUT_DOCX}")
print(f"  {len(QA)} questions | {yes_count} Yes | {partial_count} Partial")
