"""
Export all RFP sheets as a single Q&A Word document.
Reads questions from the original RFP and answers from the filled output.

Sheet column mapping (from inspection):
  AI / Data / Functional  — question=col5, yes=col6, resp=col7, cite=col8  (rows 6+)
  Supplier Questions      — question=col4 (original), resp=col4, cite=col5 (filled)  (rows 6+)
  Technical Questions     — question=col5, resp=col6, cite=col7  (no compliance col)  (rows 6+)
"""
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_RFP    = "/Users/Adam.Boyle/Downloads/Attachment 2_Ralph Lauren Contact Center Technology RFP Vendor Requirements_Questionnaire_Pricing.xlsx"
INPUT_FILLED = "/Users/Adam.Boyle/Downloads/NiCE Response - Ralph Lauren CCaaS RFP.xlsx"
OUTPUT_DOCX  = "/Users/Adam.Boyle/Downloads/NiCE - Ralph Lauren Full RFP Response.docx"

NICE_NAVY   = RGBColor(0x00, 0x29, 0x5E)
NICE_TEAL   = RGBColor(0x00, 0x97, 0x9D)
AMBER       = RGBColor(0xC4, 0x73, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY        = RGBColor(0x88, 0x88, 0x88)
DARK_TEXT   = RGBColor(0x22, 0x22, 0x22)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

wb_rfp    = openpyxl.load_workbook(INPUT_RFP)
wb_filled = openpyxl.load_workbook(INPUT_FILLED)

# ── Sheet definitions ──────────────────────────────────────────────────────────
# Each entry: (display_name, original_sheet_key, filled_sheet_key, schema)
# schema: 'assessment' | 'supplier' | 'technical'
SHEETS = [
    {
        "title": "AI & Virtual Agent Assessment",
        "orig_key": "2.AI & Virtual Agent Assessment",
        "fill_key": "2.AI & Virtual Agent Assessment",
        "schema": "assessment",
    },
    {
        "title": "Data Assessment",
        "orig_key": "3. Data Assessment",
        "fill_key": "3. Data Assessment",
        "schema": "assessment",
    },
    {
        "title": "Functional Assessment",
        "orig_key": "4. Functional Assessment",
        "fill_key": "4. Functional Assessment",
        "schema": "assessment",
    },
    {
        "title": "Supplier Questions",
        "orig_key": "5. Supplier Questions",
        "fill_key": "5. Supplier Questions",
        "schema": "supplier",
    },
    {
        "title": "Technical Questions",
        "orig_key": "6. Technical Questions",
        "fill_key": "6. Technical Questions",
        "schema": "technical",
    },
]

def load_qa(sheet_def):
    """Extract Q&A pairs for a sheet based on its schema."""
    orig_name = next((s for s in wb_rfp.sheetnames if sheet_def["orig_key"] in s or s == sheet_def["orig_key"]), None)
    fill_name = next((s for s in wb_filled.sheetnames if sheet_def["fill_key"] in s or s == sheet_def["fill_key"]), None)
    if not orig_name or not fill_name:
        print(f"  !! Could not find sheet: {sheet_def['title']}")
        return []

    ws_q = wb_rfp[orig_name]
    ws_a = wb_filled[fill_name]
    pairs = []
    schema = sheet_def["schema"]

    for row in range(6, 300):
        req_id = ws_q.cell(row=row, column=2).value
        subsec = ws_q.cell(row=row, column=3).value

        if schema == "assessment":
            cap      = ws_q.cell(row=row, column=4).value
            question = ws_q.cell(row=row, column=5).value
            if not question:
                break
            compliance = ws_a.cell(row=row, column=6).value or ""
            response   = ws_a.cell(row=row, column=7).value or ""
            citation   = ws_a.cell(row=row, column=8).value or ""

        elif schema == "supplier":
            cap      = None
            question = ws_q.cell(row=row, column=4).value   # original has question here
            if not question:
                if not req_id:
                    break
                continue
            compliance = None
            response   = ws_a.cell(row=row, column=4).value or ""  # filled overwrote col4
            citation   = ws_a.cell(row=row, column=5).value or ""

        elif schema == "technical":
            cap      = ws_q.cell(row=row, column=4).value
            question = ws_q.cell(row=row, column=5).value
            if not question:
                break
            compliance = None
            response   = ws_a.cell(row=row, column=6).value or ""
            citation   = ws_a.cell(row=row, column=7).value or ""

        else:
            break

        pairs.append({
            "id":         str(req_id or "").strip(),
            "subsec":     str(subsec or "").strip(),
            "cap":        str(cap or "").strip() if cap else "",
            "question":   str(question).strip(),
            "compliance": str(compliance).strip() if compliance else "",
            "response":   str(response).strip(),
            "citation":   str(citation).strip(),
        })

    return pairs

# ── Build Word document ───────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# Cover title
title = doc.add_heading("Ralph Lauren CCaaS RFP", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = NICE_NAVY
    run.font.size = Pt(22)

sub = doc.add_paragraph("NiCE Response — Full Questionnaire")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = NICE_TEAL
    run.font.size = Pt(12)
    run.font.bold = True

doc.add_paragraph()

# Process each sheet
total_questions = 0

for sheet_def in SHEETS:
    pairs = load_qa(sheet_def)
    if not pairs:
        continue

    total_questions += len(pairs)
    print(f"  {sheet_def['title']}: {len(pairs)} Q&As")

    # ── Sheet-level heading (page break for all but the first) ─────────────
    # Add page break before each sheet section (except the first)
    section_heading = doc.add_heading(sheet_def["title"], level=1)
    for run in section_heading.runs:
        run.font.color.rgb = WHITE
        run.font.size = Pt(14)
        run.font.bold = True
    # Navy background for section banner
    shd_elem = OxmlElement('w:shd')
    shd_elem.set(qn('w:val'), 'clear')
    shd_elem.set(qn('w:color'), 'auto')
    shd_elem.set(qn('w:fill'), '00295E')
    pPr = section_heading._p.get_or_add_pPr()
    pPr.append(shd_elem)
    section_heading.paragraph_format.space_before = Pt(20)
    section_heading.paragraph_format.space_after  = Pt(8)

    # Stats line (Yes / Partial count)
    yes_count     = sum(1 for p in pairs if p["compliance"] == "Yes")
    partial_count = sum(1 for p in pairs if p["compliance"] == "Partial")
    if yes_count or partial_count:
        stats = doc.add_paragraph()
        stats.paragraph_format.space_before = Pt(0)
        stats.paragraph_format.space_after  = Pt(10)
        stats.add_run(f"{len(pairs)} requirements   ").font.size = Pt(9)
        yes_run = stats.add_run(f"✓ {yes_count} Yes   ")
        yes_run.font.size = Pt(9); yes_run.font.bold = True; yes_run.font.color.rgb = NICE_NAVY
        if partial_count:
            par_run = stats.add_run(f"◑ {partial_count} Partial")
            par_run.font.size = Pt(9); par_run.font.bold = True; par_run.font.color.rgb = AMBER

    # ── Q&A entries ────────────────────────────────────────────────────────
    current_subsec = None
    for i, qa in enumerate(pairs, 1):

        # Subsection heading
        if qa["subsec"] and qa["subsec"] != current_subsec:
            current_subsec = qa["subsec"]
            h = doc.add_heading(current_subsec, level=2)
            for run in h.runs:
                run.font.color.rgb = NICE_NAVY
                run.font.size = Pt(11)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(3)

        # ID + capability label
        label_para = doc.add_paragraph()
        label_para.paragraph_format.space_before = Pt(8)
        label_para.paragraph_format.space_after  = Pt(2)
        if qa["id"]:
            id_run = label_para.add_run(f"{qa['id']}  ")
            id_run.font.bold = True
            id_run.font.color.rgb = NICE_TEAL
            id_run.font.size = Pt(8.5)
        if qa["cap"]:
            cap_run = label_para.add_run(qa["cap"])
            cap_run.font.bold = True
            cap_run.font.color.rgb = NICE_NAVY
            cap_run.font.size = Pt(9.5)

        # Question text
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(1)
        q_para.paragraph_format.space_after  = Pt(4)
        q_para.paragraph_format.left_indent  = Inches(0.1)
        q_run = q_para.add_run(qa["question"])
        q_run.font.size = Pt(9.5)
        q_run.font.color.rgb = DARK_TEXT

        # Answer row
        if qa["response"]:
            tbl = doc.add_table(rows=1, cols=2 if qa["compliance"] else 1)
            tbl.style = "Table Grid"
            tbl.allow_autofit = False

            if qa["compliance"]:
                badge_cell = tbl.cell(0, 0)
                badge_cell.width = Inches(0.7)
                badge_para = badge_cell.paragraphs[0]
                badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                badge_run = badge_para.add_run(qa["compliance"])
                badge_run.font.bold = True
                badge_run.font.size = Pt(9)
                badge_run.font.color.rgb = WHITE
                if qa["compliance"] == "Yes":
                    set_cell_bg(badge_cell, "00295E")
                elif qa["compliance"] == "Partial":
                    set_cell_bg(badge_cell, "C47300")
                else:
                    set_cell_bg(badge_cell, "C41E1E")
                resp_cell = tbl.cell(0, 1)
                tbl.columns[0].width = Inches(0.7)
                tbl.columns[1].width = Inches(5.4)
            else:
                resp_cell = tbl.cell(0, 0)
                tbl.columns[0].width = Inches(6.1)

            resp_para = resp_cell.paragraphs[0]
            resp_run  = resp_para.add_run(qa["response"])
            resp_run.font.size = Pt(9.5)

            if qa["citation"]:
                cite_para = resp_cell.add_paragraph()
                cite_run = cite_para.add_run(f"↗  {qa['citation']}")
                cite_run.font.size = Pt(8)
                cite_run.font.color.rgb = NICE_TEAL
                cite_run.font.italic = True

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(1)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph("Prepared by NICE  |  Confidential — Ralph Lauren CCaaS RFP Response  |  April 2026")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.italic = True

doc.save(OUTPUT_DOCX)
print(f"\n✓ Saved: {OUTPUT_DOCX}")
print(f"  {total_questions} total Q&As across {len(SHEETS)} sheets")
