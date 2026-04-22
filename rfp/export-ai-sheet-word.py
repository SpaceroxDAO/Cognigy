"""
Export AI & Virtual Agent Assessment sheet as a Q&A Word document.
Reads questions from the original RFP and answers from the filled output.
"""
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_RFP    = "/Users/Adam.Boyle/Downloads/Attachment 2_Ralph Lauren Contact Center Technology RFP Vendor Requirements_Questionnaire_Pricing.xlsx"
INPUT_FILLED = "/Users/Adam.Boyle/Downloads/NiCE Response - Ralph Lauren CCaaS RFP.xlsx"
OUTPUT_DOCX  = "/Users/Adam.Boyle/Downloads/NiCE - Ralph Lauren AI Assessment Q&A.docx"

NICE_NAVY   = RGBColor(0x00, 0x29, 0x5E)   # NICE brand navy
NICE_TEAL   = RGBColor(0x00, 0x97, 0x9D)   # NICE brand teal
PARTIAL_COL = RGBColor(0xC4, 0x73, 0x00)   # amber for Partial
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Load data ─────────────────────────────────────────────────────────────────
wb_rfp    = openpyxl.load_workbook(INPUT_RFP)
wb_filled = openpyxl.load_workbook(INPUT_FILLED)

ai_sheet_name = next(s for s in wb_rfp.sheetnames if 'AI' in s and 'Virtual' in s)
ws_q = wb_rfp[ai_sheet_name]

ai_filled_name = next(s for s in wb_filled.sheetnames if 'AI' in s and 'Virtual' in s)
ws_a = wb_filled[ai_filled_name]

# Extract Q&A pairs (rows 6 onward, until blank requirement)
qa_pairs = []
for row in range(6, 200):
    req_id    = ws_q.cell(row=row, column=2).value   # ID
    subsec    = ws_q.cell(row=row, column=3).value   # Subsection
    cap       = ws_q.cell(row=row, column=4).value   # Capability
    question  = ws_q.cell(row=row, column=5).value   # Requirements
    if not question:
        break
    compliance = ws_a.cell(row=row, column=6).value or ""
    response   = ws_a.cell(row=row, column=7).value or ""
    citation   = ws_a.cell(row=row, column=8).value or ""
    qa_pairs.append({
        "id": str(req_id or ""),
        "subsec": str(subsec or ""),
        "cap": str(cap or ""),
        "question": str(question).strip(),
        "compliance": str(compliance).strip(),
        "response": str(response).strip(),
        "citation": str(citation).strip(),
    })

print(f"Loaded {len(qa_pairs)} Q&A pairs")

# ── Build Word document ───────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# Title
title = doc.add_heading("AI & Virtual Agent Assessment", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = NICE_NAVY
    run.font.size = Pt(20)

# Subtitle
sub = doc.add_paragraph("NiCE Response — Ralph Lauren CCaaS RFP")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = NICE_TEAL
    run.font.size = Pt(11)
    run.font.bold = True

doc.add_paragraph()  # spacer

# Group by subsection for section headings
current_subsec = None
for i, qa in enumerate(qa_pairs, 1):
    # Section heading on subsection change
    if qa["subsec"] != current_subsec:
        current_subsec = qa["subsec"]
        heading = doc.add_heading(current_subsec, level=1)
        for run in heading.runs:
            run.font.color.rgb = NICE_NAVY
            run.font.size = Pt(13)
        heading.paragraph_format.space_before = Pt(16)
        heading.paragraph_format.space_after  = Pt(4)

    # ── Question block ─────────────────────────────────────────────────────
    # ID + capability label
    label_para = doc.add_paragraph()
    label_para.paragraph_format.space_before = Pt(10)
    label_para.paragraph_format.space_after  = Pt(2)
    id_run = label_para.add_run(f"{qa['id']}  ")
    id_run.font.bold = True
    id_run.font.color.rgb = NICE_TEAL
    id_run.font.size = Pt(9)
    cap_run = label_para.add_run(qa["cap"])
    cap_run.font.bold = True
    cap_run.font.color.rgb = NICE_NAVY
    cap_run.font.size = Pt(10)

    # Question text
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_before = Pt(1)
    q_para.paragraph_format.space_after  = Pt(4)
    q_para.paragraph_format.left_indent  = Inches(0.15)
    q_run = q_para.add_run(qa["question"])
    q_run.font.size = Pt(10)
    q_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Answer block (table for badge + text) ─────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.allow_autofit = False

    # Badge cell (Yes / Partial / No)
    badge_cell = tbl.cell(0, 0)
    badge_cell.width = Inches(0.75)
    badge_para = badge_cell.paragraphs[0]
    badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_run = badge_para.add_run(qa["compliance"])
    badge_run.font.bold = True
    badge_run.font.size = Pt(10)
    badge_run.font.color.rgb = WHITE

    if qa["compliance"] == "Yes":
        set_cell_bg(badge_cell, "00295E")   # navy
    elif qa["compliance"] == "Partial":
        set_cell_bg(badge_cell, "C47300")   # amber
    else:
        set_cell_bg(badge_cell, "C41E1E")   # red

    # Response cell
    resp_cell = tbl.cell(0, 1)
    resp_para = resp_cell.paragraphs[0]
    resp_run = resp_para.add_run(qa["response"])
    resp_run.font.size = Pt(10)

    # Citation below response (if present)
    if qa["citation"]:
        cite_para = resp_cell.add_paragraph()
        cite_run = cite_para.add_run(f"↗  {qa['citation']}")
        cite_run.font.size = Pt(8)
        cite_run.font.color.rgb = NICE_TEAL
        cite_run.font.italic = True

    # Column widths
    tbl.columns[0].width = Inches(0.75)
    tbl.columns[1].width = Inches(5.2)

    tbl.rows[0].height = None  # auto height

    # Spacer after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(2)

# Footer note
doc.add_paragraph()
footer = doc.add_paragraph("Prepared by NICE  |  Confidential — Ralph Lauren CCaaS RFP Response")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True

doc.save(OUTPUT_DOCX)
print(f"\n✓ Saved: {OUTPUT_DOCX}")
print(f"  {len(qa_pairs)} questions | {sum(1 for q in qa_pairs if q['compliance']=='Yes')} Yes | {sum(1 for q in qa_pairs if q['compliance']=='Partial')} Partial")
