from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- Styles helper ---
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bold_intro(label, rest):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Light grey background via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3A5F')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            for run in row_cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()  # space after table

def add_callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8DC')
    pPr.append(shd)
    return p

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Cognigy Demo Generation with Claude Code', 0)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
title.runs[0].font.size = Pt(24)

add_bold_intro('By ', 'Adam Boyle | NiCE + Cognigy Professional Services')
doc.add_paragraph()

add_body(
    'This page documents how to use Claude Enterprise to power Claude Code — Anthropic\'s AI-native CLI — '
    'as a fully autonomous Cognigy demo package generator. Rather than building demos by hand inside Cognigy.AI\'s UI, '
    'this approach lets Claude Code generate production-ready, importable .zip packages from a prospect name and a few minutes of setup.'
)

# ============================================================
# WHAT IS CLAUDE CODE
# ============================================================
add_heading('What Is Claude Code?', level=1)
add_body(
    'Claude Code is a CLI tool (and IDE extension) that gives Claude direct access to your local filesystem, terminal, '
    'and codebase. Unlike Claude.ai in a browser, Claude Code can read, write, and execute files — making it capable of '
    'acting as a development agent, not just a chat assistant.'
)
add_body(
    'In this use case, Claude Code is the engine behind a Cognigy Demo Generation Agent: a dedicated local project that '
    'holds all the knowledge, patterns, and build infrastructure needed to go from a prospect name to a ready-to-import '
    'Cognigy package.'
)

# ============================================================
# USE CASE
# ============================================================
add_heading('Use Case: Automated Cognigy Demo Package Generation', level=1)

add_heading('Overview', level=2)
add_body(
    'Building a Cognigy demo from scratch — configuring flows, code nodes, AI Agent jobs, tools, voice settings, and all '
    'the required JSON cross-references — is time-consuming and error-prone. A single demo can involve dozens of '
    'interdependent files that must be consistent with each other to import successfully.'
)
add_body(
    'This use case encodes that entire process into a Claude Code project, backed by a set of Skills (custom instruction '
    'files) that Claude loads on demand. The result: a consultant can describe a prospect and get a fully functional demo '
    'package — complete with realistic tools, branded voice persona, and correct Cognigy package structure — without '
    'touching the Cognigy UI until it\'s time to import.'
)

add_bold_intro('What Claude Code generates:', '')
doc.add_paragraph()
add_bullet('A .zip package ready to import into any Cognigy.AI environment')
add_bullet('A complete AI Agent flow with tools, code nodes, and voice configuration')
add_bullet('Prospect-specific FAQ knowledge, instructions, and mock data')
add_bullet('Optional xApp (mobile web) components surfaced during the voice conversation')

# ============================================================
# SETUP
# ============================================================
add_heading('Setup: Creating Your Demo Generation Directory', level=1)

add_body(
    'The best practice is to keep the Cognigy demo generation project in its own dedicated directory, separate from any '
    'client project files. This directory becomes Claude Code\'s persistent workspace — it holds the generator code, the '
    'reference package template, all Skills, and every build script.'
)

add_heading('Recommended Directory Structure', level=2)

add_code_block([
    '~/Cognigy/                          ← Dedicated root for the agent',
    '  CLAUDE.md                         ← Claude Code project instructions (always in context)',
    '  cognigy-package-generator/',
    '    clone-and-modify.js             ← Core generator: clones template, swaps content',
    '    index.js                        ← Legacy (do not use)',
    '    lib/                            ← Legacy sub-modules (reference only)',
    '  credit-card-analysis/             ← Reference template: extracted working Cognigy package',
    '  .claude/',
    '    skills/                         ← Skill files loaded on demand by Claude Code',
    '      demo-builder/',
    '      prospect-research/',
    '      tool-design/',
    '      instruction-patterns/',
    '      voice-config/',
    '      package-builder/',
    '      cognigy-package-format/',
    '      demo-review/',
    '      node-reference/',
    '      rfp-responder/',
    '    memory/                         ← Auto-memory: persists context across sessions',
    '      MEMORY.md                     ← Index of all saved memories',
    '      *.md                          ← Individual memory files by topic',
    '  build-*.js                        ← Per-demo build scripts (one per prospect)',
    '  output/                           ← Generated .zip packages',
])

doc.add_paragraph()
add_heading('Why a Dedicated Directory?', level=2)
add_bullet(
    'Claude Code loads CLAUDE.md automatically — keeping it in its own directory means it always has the right '
    'context without contaminating other projects.'
)
add_bullet(
    'The memory/ folder persists knowledge across sessions: lessons learned, patterns confirmed, demo history, '
    'and user preferences — so you don\'t re-explain your conventions every conversation.'
)
add_bullet(
    'Build scripts accumulate over time and become a library of working examples for future demos.'
)

# ============================================================
# SETUP STEPS
# ============================================================
add_heading('How to Set Up Claude Code', level=1)

add_heading('1. Install Claude Code', level=2)
add_code_block(['npm install -g @anthropic-ai/claude-code'])
doc.add_paragraph()
add_body('Claude Code is also available as a VS Code extension and a JetBrains plugin.')

add_heading('2. Connect to Claude Enterprise', level=2)
add_body('Launch Claude Code from your demo directory:')
add_code_block(['cd ~/Cognigy', 'claude'])
doc.add_paragraph()
add_body(
    'On first launch, authenticate with your Claude Enterprise account. Your organization\'s enterprise subscription '
    'powers all Claude Code sessions — no separate API key required.'
)

add_heading('3. Initialize the Project', level=2)
add_body(
    'Drop a CLAUDE.md file in your root directory. This file is the agent\'s brain — it describes:'
)
add_bullet('What the project does (generate Cognigy demo packages)')
add_bullet('Which method to use (clone-and-modify, not from scratch)')
add_bullet('Critical package requirements (ZIP format constraints, required node fields)')
add_bullet('Standard flow patterns (Start → Once → AI Agent Job → tools)')
add_bullet('When to load which Skill (trigger table mapping task types to skill files)')
add_body('Claude Code reads CLAUDE.md at the start of every session. Everything in it overrides Claude\'s default behavior.')

add_heading('4. Add Your Skills', level=2)
add_body(
    'Skills are Markdown files saved in .claude/skills/. Each skill encodes a specialized body of knowledge — '
    'how to design tools, how to configure voice settings, how to structure instructions, etc. '
    'Claude Code loads them on demand based on the task at hand, keeping each session focused.'
)

# ============================================================
# HOW IT WORKS
# ============================================================
add_heading('How It Works: The Demo Generation Flow', level=1)

add_heading('Step 1: Name the Prospect', level=2)
add_body('Open Claude Code in the ~/Cognigy directory and run:')
add_code_block(['/create-demo'])
doc.add_paragraph()
add_body(
    'Claude Code asks for the prospect name, industry, and use case. It then automatically triggers the '
    'demo-builder skill, which orchestrates the remaining steps.'
)

add_heading('Step 2: Research (Automated)', level=2)
add_body(
    'The prospect-research skill kicks in. Claude Code uses web search and page fetching to find:'
)
add_bullet('The prospect\'s real FAQ content, product names, and support policies')
add_bullet('IVR pain points and call routing context')
add_bullet('Terminology, tone, and brand voice')
add_body('This produces a Scraped Knowledge Summary for review before anything is built.')

add_heading('Step 3: Spec Design', level=2)
add_body('Claude Code drafts a build spec:')
add_bullet('Tools: what the AI Agent can do (check status, update records, schedule callbacks, send xApp links, etc.)')
add_bullet('Instructions: the context.instructions prompt with standard voice rules, confirmation patterns, and domain-specific guardrails')
add_bullet('Knowledge: the context.knowledge block with FAQ content from real sources')
add_bullet('Voice config: STT/TTS settings, pronunciation rules, silence handling')

add_heading('Step 4: Build', level=2)
add_body('Claude Code writes a build-[prospect].js script and runs it:')
add_code_block(['node build-benesys.js'])
doc.add_paragraph()
add_body(
    'The script calls cloneAndModify() — a function that takes the reference package, rebuilds all JSON files with the '
    'new content, and produces a valid .zip file. Every cross-reference (flow ↔ chart ↔ nodeData ↔ locale), every '
    'required field, and every mock code node is handled automatically.'
)

add_heading('Step 5: Import and Test', level=2)
add_body(
    'Import the .zip into Cognigy.AI. Run the demo. If anything needs adjusting, describe it to Claude Code and it '
    'edits the build script and regenerates — no manual JSON editing required.'
)

# ============================================================
# SKILLS REFERENCE
# ============================================================
add_heading('Skills Reference', level=1)
add_body(
    'Skills are the core of what makes Claude Code consistent. Each one encodes standards that would otherwise '
    'require tribal knowledge or careful prompting.'
)
doc.add_paragraph()

add_table(
    ['Skill', 'What It Encodes'],
    [
        ['demo-builder', 'End-to-end orchestration: research → spec → build → validate → deliver'],
        ['package-builder', 'How clone-and-modify.js works; ZIP format constraints; failure modes'],
        ['tool-design', 'Tool schemas, mock data patterns, answer templates, common mistakes'],
        ['instruction-patterns', 'Standard voice rules (mandatory for every demo), prompt structure, three code node chain'],
        ['voice-config', 'TTS/STT settings, pronunciation rules, barge-in, DTMF, silence handling'],
        ['prospect-research', 'How to scrape real FAQ content and structure it for context.knowledge'],
        ['cognigy-package-format', 'ZIP structure, JSON schemas, cross-reference rules — the debug reference'],
        ['demo-review', 'QA from pasted transcripts or session logs; severity guide; output format'],
        ['node-reference', 'Every Cognigy node type\'s config fields and extension name'],
        ['rfp-responder', 'Fill RFP/RFI Excel questionnaires using Python + openpyxl'],
    ]
)

add_body(
    'Skills are not loaded all at once — CLAUDE.md contains a trigger table that tells Claude Code which skills '
    'to load based on the task. This keeps sessions focused and prevents instructions from different domains from conflicting.'
)

# ============================================================
# MEMORY
# ============================================================
add_heading('Memory: Persistence Across Sessions', level=1)
add_body(
    'Claude Code maintains a memory/ folder that persists facts across sessions — things that would otherwise need '
    'to be re-explained every time:'
)
add_bullet('Project memory: the 14+ demos built, patterns that were validated, decisions made')
add_bullet('Feedback memory: what to avoid (e.g., never write required: [] in a tool schema — omit the key), what to keep doing')
add_bullet('User memory: the demo persona, voice stack, lab environment')
add_bullet('Reference memory: where to find things (Cognigy CLI notes, RFP outputs, etc.)')
add_body(
    'Memory is stored as individual .md files, indexed in MEMORY.md. Claude Code reads the index at session start '
    'and pulls relevant files when they\'re needed.'
)

# ============================================================
# WHY IT WORKS
# ============================================================
add_heading('Why This Approach Works', level=1)

add_bold_intro('Consistency: ', 'Every demo follows the same flow pattern, uses the same voice rules, and passes the same import requirements — regardless of how the task is phrased.')
add_bold_intro('Speed: ', 'A demo that previously took a day of manual UI work can be built in under an hour, including research.')
add_bold_intro('Iteration: ', 'Changing a tool, rewriting the instructions, or adding an xApp component takes seconds — edit the build script, rerun, reimport.')
add_bold_intro('Institutional knowledge capture: ', 'Skills and memory mean the patterns developed across 14+ demos are encoded and reusable. New team members get the same output quality immediately.')
add_bold_intro('No hallucination on structure: ', 'Claude Code doesn\'t guess at the Cognigy package format — it clones a verified working package and modifies it. The structural correctness is inherited, not generated.')

# ============================================================
# GETTING STARTED
# ============================================================
add_heading('Getting Started', level=1)

add_bullet('Install Claude Code: npm install -g @anthropic-ai/claude-code')
add_bullet('Create your ~/Cognigy directory and add CLAUDE.md and the Skills folder')
add_bullet('Clone a working Cognigy package to use as the reference template')
add_bullet('Run claude from the directory and type /create-demo')

doc.add_paragraph()
add_body('For questions about this setup, the Skills library, or access to the reference package and build scripts, reach out to Adam Boyle.')

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Last updated: April 2026')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
output_path = '/Users/Adam.Boyle/Cognigy/Cognigy-Demo-Generation-Claude-Code.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
