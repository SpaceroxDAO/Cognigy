"""
Generate ComEd IVA FAQ Test Plan as a Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/Users/Adam.Boyle/Cognigy/ComEd-IVA-Test-Plan.docx"

# ── Brand colours ────────────────────────────────────────────────────
COMED_BLUE  = RGBColor(0x00, 0x3B, 0x6E)   # ComEd dark navy
COMED_RED   = RGBColor(0xC8, 0x10, 0x26)   # ComEd red
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
MID_GREY    = RGBColor(0x60, 0x60, 0x60)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(14)
    r.font.color.rgb = COMED_BLUE
    # Red underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C81026")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(11)
    r.font.color.rgb = COMED_RED
    return p

def body(text, space_after=4, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, text, italic=italic, color=color)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True)
        add_run(p, text)
    else:
        add_run(p, text)
    return p

def expect_box(text):
    """Shaded expect block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Inches(0.25)
    add_run(p, "What to expect:  ", bold=True, color=COMED_BLUE)
    add_run(p, text, color=MID_GREY)
    return p

def make_table(headers, rows, header_bg="003B6E"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return t

# ═══════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(6)
title_p.paragraph_format.space_after  = Pt(2)
r = title_p.add_run("ComEd Virtual Assistant")
r.bold = True
r.font.size  = Pt(22)
r.font.color.rgb = COMED_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
r2 = sub_p.add_run("Test Guide")
r2.bold = True
r2.font.size  = Pt(16)
r2.font.color.rgb = COMED_RED

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(16)
add_run(meta_p, "April 2026  |  Powered by NiCE + Cognigy.AI", color=MID_GREY, size=9)

# Divider
div = doc.add_paragraph()
div.paragraph_format.space_after = Pt(14)
pPr = div._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "12")
bot.set(qn("w:space"), "1")
bot.set(qn("w:color"), "003B6E")
pBdr.append(bot)
pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════
# SECTION 1 — HOW TO TEST
# ═══════════════════════════════════════════════════════════════════════
heading1("How to Test")
body(
    "Access the assistant using the web link or phone number provided by your NiCE contact. "
    "Simply have a conversation the way a real customer would. "
    "The questions below are starting points — feel free to rephrase them naturally."
)

# ═══════════════════════════════════════════════════════════════════════
# SECTION 2 — TOPICS
# ═══════════════════════════════════════════════════════════════════════
heading1("What to Test By Topic")

# --- Power Outages ---
heading2("Power Outages")
body("Try asking things like:")
for q in [
    "I don't have power",
    "My lights just went out",
    "How do I report an outage?",
    "Is there an outage in my area?",
    "What do I do if I see a downed power line?",
    "How do I get outage text alerts?",
]:
    bullet(q)
expect_box(
    "Jane should tell you how to report the outage, give you the phone number (1-800-EDISON-1), "
    "mention texting OUT to 26633, and reference the ComEd app and website. "
    "For downed lines, she should lead with safety — stay back, call 911 first. "
    "If you ask for a link, she should text one to your phone within a few seconds."
)

# --- Billing & Payments ---
heading2("Billing & Payments")
body("Try asking things like:")
for q in [
    "How do I pay my bill?",
    "What if I can't afford my bill this month?",
    "How do I sign up for AutoPay?",
    "What is Budget Billing?",
    "How do I download a copy of my bill?",
    "Can you send me the link to set up automatic payments?",
]:
    bullet(q)
expect_box(
    "Jane should explain the available payment options (online, by mail, in person, AutoPay, E-Check). "
    "For customers who can't pay, she should mention payment arrangements and assistance programs — "
    "not make promises about specific amounts. If you ask for a link, she should text you the relevant page."
)

# --- Start / Stop / Move ---
heading2("Starting, Stopping, or Moving Service")
body("Try asking things like:")
for q in [
    "How do I start electric service?",
    "I'm moving next month — how do I transfer my service?",
    "How do I stop my service?",
    "I'm a business owner, how do I set up service?",
    "Why is Experian asking me questions when I try to start service?",
]:
    bullet(q)
expect_box(
    "Jane should walk you through the process — what information you'll need, whether to go online or call, "
    "and the relevant phone numbers. She should never ask you to provide personal information "
    "like your Social Security number during the call."
)

# --- My Account ---
heading2("My Account & Online Access")
body("Try asking things like:")
for q in [
    "How do I create an online account?",
    "I can't log in to my account",
    "How do I update my phone number?",
    "What is Two-Step Verification?",
    "How do I change my electricity supplier?",
    "Can you text me the link to My Account?",
]:
    bullet(q)
expect_box(
    "Jane should explain the steps clearly in plain language — what information you need to register, "
    "how to reset your login, how to navigate to the right settings page. "
    "For anything requiring a link, she texts it to you directly."
)

# --- General ---
heading2("General Questions")
body("Try asking things like:")
for q in [
    "How do I read my electric meter?",
    "How do I submit a claim for food that spoiled during an outage?",
    "What browsers does ComEd's website support?",
]:
    bullet(q)
expect_box(
    "Jane should answer directly and conversationally, without making you repeat yourself "
    "or navigate a menu."
)

# ═══════════════════════════════════════════════════════════════════════
# SECTION 3 — MULTILINGUAL
# ═══════════════════════════════════════════════════════════════════════
heading1("Testing in Other Languages")
body(
    "Jane is designed to detect your language automatically and respond in kind. "
    "Try starting a conversation in any of the languages below — she should switch immediately without being asked."
)

make_table(
    ["Language", "Sample Opening", "Translation"],
    [
        ("Spanish",   '"¿Cómo pago mi factura?"',                    "How do I pay my bill?"),
        ("Polish",    '"Jak zgłosić awarię prądu?"',                 "How do I report a power outage?"),
        ("Hindi",     '"मेरे घर में बिजली नहीं है"',                "I don't have electricity at home"),
        ("Mandarin",  '"我怎么查看我的账单？"',                       "How do I check my bill?"),
        ("Urdu",      '"میں اپنا اکاؤنٹ کیسے بنا سکتا ہوں؟"',      "How do I create my account?"),
        ("Kannada",   '"ನನ್ನ ಸೇವೆ ಹೇಗೆ ಪ್ರಾರಂಭಿಸುವುದು?"',          "How do I start my service?"),
        ("Ukrainian", '"У мене немає електрики"',                    "I don't have electricity"),
        ("Russian",   '"Как мне оплатить счёт?"',                    "How do I pay my bill?"),
    ]
)

expect_box(
    "Jane should respond fully in the language you used — not switch back to English — "
    "and give the same quality of answer she would in English."
)

# ═══════════════════════════════════════════════════════════════════════
# SECTION 4 — ALWAYS / NEVER
# ═══════════════════════════════════════════════════════════════════════
heading1("What Jane Should Always Do")
for item in [
    "Keep answers short and conversational — one or two sentences at a time",
    "Offer to text you a link rather than reading out a web address",
    "Follow your lead on language for the whole conversation",
    "Offer to connect you to a live agent if she can't help",
]:
    bullet(item)

doc.add_paragraph()

heading1("What Jane Should Never Do")
for item in [
    "Make up information or guess at policy details",
    "Read out a full web URL",
    "Ask you to provide your Social Security number or full account details",
    "Promise a specific time for power restoration",
    "Give your account balance or bill history (she does not have access to individual accounts in this test environment)",
]:
    bullet(item)

# ═══════════════════════════════════════════════════════════════════════
# SECTION 5 — REPORTING
# ═══════════════════════════════════════════════════════════════════════
heading1("Reporting Feedback")
body("If something doesn't feel right — wrong answer, wrong language, unexpected behavior — note down:")
bullet("What you asked (as close to word-for-word as you remember)")
bullet("What Jane said back")
bullet("What you expected her to say")

doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, "Send feedback to:  ", bold=True)
add_run(p, "[ComEd project contact — provided by your NiCE SE]", italic=True, color=MID_GREY)

# ── Save ─────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print("Saved: " + OUTPUT)
